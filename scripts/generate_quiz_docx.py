#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

QUIZ = [
    {
        "q": "Welcher Modifikator ist am restriktivsten?",
        "opts": ["public", "protected", "(kein Modifikator) package-private", "private"],
        "multi": False
    },
    {
        "q": "Wofür steht „package-private“?",
        "opts": [
            "Sichtbar in derselben Klasse und in Subklassen anderer Pakete",
            "Nur im selben Paket sichtbar",
            "Überall sichtbar",
            "Nur innerhalb der Klasse sichtbar"
        ],
        "multi": False
    },
    {
        "q": "Welche Aussage ist korrekt? (Mehrfachauswahl möglich)",
        "opts": [
            "Attribute sollten in der Regel private sein",
            "Getter dürfen niemals verwendet werden",
            "Setter können Validierung enthalten",
            "Kapselung erhöht die Wartbarkeit des Codes"
        ],
        "multi": True
    },
    {
        "q": "Welche Rückgabe ist bei Sammlungen aus Kapselungssicht vorzuziehen?",
        "opts": [
            "Die interne, veränderbare Liste",
            "Eine unveränderliche Sicht (unmodifiable)",
            "Immer null",
            "Eine defensive Kopie"
        ],
        "multi": True
    },
    {
        "q": "Wofür eignet sich protected?",
        "opts": [
            "Für API-Methoden, die von jedem aufgerufen werden sollen",
            "Für Methoden, die nur Subklassen (und Paket-Mitglieder) sehen dürfen",
            "Für komplett interne Hilfsmethoden",
            "Für konstante Werte"
        ],
        "multi": False
    }
]


def add_checkbox_run(p):
    r = p.add_run("☐ ")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    return r


def main():
    doc = Document()

    # Titel
    title = doc.add_paragraph()
    title_run = title.add_run("Kapselung – Quiz")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta
    meta = doc.add_paragraph()
    meta.add_run("Name: ").bold = True
    meta.add_run("____________________________    ")
    meta.add_run("Datum: ").bold = True
    meta.add_run("____________________________")

    doc.add_paragraph("")

    # Hinweis
    hint = doc.add_paragraph("Hinweis: Bei Mehrfachauswahl-Fragen können mehrere Kästchen anzukreuzen sein.")
    hint.runs[0].italic = True

    # Fragen
    for idx, item in enumerate(QUIZ, start=1):
        doc.add_paragraph("")
        q = doc.add_paragraph(f"{idx}. {item['q']}")
        for opt in item["opts"]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_checkbox_run(p)
            p.add_run(opt)

    # Platz für Bemerkungen
    doc.add_paragraph("")
    doc.add_paragraph("Bemerkungen:")
    for _ in range(3):
        doc.add_paragraph("_______________________________________________")

    doc.save("materials/quiz_kapselung.docx")


if __name__ == "__main__":
    main()
